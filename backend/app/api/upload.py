"""
文档上传解析端点 —— 支持 DOCX、PDF、Markdown、图片等常见文档格式。

接受 multipart/form-data 文件上传，解析为纯文本后返回。
图片文件通过 VisionPreprocessorTool（Qwen VL）转换为 Markdown。
"""

import io
import os
from fastapi import APIRouter, File, HTTPException, UploadFile
from pydantic import BaseModel, Field

router = APIRouter(prefix="/api", tags=["upload"])

# 文件大小限制：文本/文档 10MB，图片 20MB
MAX_FILE_SIZE = 10 * 1024 * 1024
MAX_IMAGE_SIZE = 20 * 1024 * 1024

# 支持的文本格式（直接 UTF-8 读取）
TEXT_EXTENSIONS = {
    ".md", ".markdown", ".txt", ".json", ".csv",
    ".html", ".htm", ".xml", ".yaml", ".yml", ".log",
    ".py", ".js", ".ts", ".jsx", ".tsx", ".css", ".sql",
}

# 需要后端解析的二进制格式
BINARY_EXTENSIONS = {".docx", ".pdf"}

# 图片格式（通过 VisionPreprocessorTool 解析）
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp"}

ALL_SUPPORTED = TEXT_EXTENSIONS | BINARY_EXTENSIONS | IMAGE_EXTENSIONS


# ----------------------------------------------------------------
# 响应模型
# ----------------------------------------------------------------

class UploadResponse(BaseModel):
    filename: str = Field(..., description="原始文件名")
    content: str = Field(..., description="解析后的纯文本内容")
    format: str = Field(..., description="文件格式")
    word_count: int = Field(default=0, description="字数统计（中文字符 + 英文单词）")

    model_config = {"extra": "forbid"}


# ----------------------------------------------------------------
# 解析器
# ----------------------------------------------------------------

def _parse_docx(file_bytes: bytes) -> str:
    """使用 python-docx 提取 Word 文档文本。"""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 同时提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            row_texts: list[str] = []
            for cell in row.cells:
                ct = cell.text.strip()
                if ct:
                    row_texts.append(ct)
            if row_texts:
                paragraphs.append(" | ".join(row_texts))

    return "\n\n".join(paragraphs)


def _parse_pdf(file_bytes: bytes) -> str:
    """使用 pdfplumber 提取 PDF 文本。"""
    import pdfplumber

    pages_text: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)

    return "\n\n".join(pages_text)


async def _parse_image(file_bytes: bytes, file_name: str) -> str:
    """通过 VisionPreprocessorTool 将图片转换为 Markdown。"""
    from app.container import get_container
    from app.tools.vision_models import VisionInput

    container = get_container()
    vision_tool = container.vision_tool

    result = await vision_tool.run(VisionInput(
        image_bytes=file_bytes,
        file_name=file_name,
    ))

    if not result.ok:
        raise RuntimeError(f"图片识别失败: {result.error}")

    return result.data.cleaned_markdown


def _detect_format(filename: str) -> str:
    """根据文件扩展名推断格式标签。"""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    mapping = {
        "docx": "docx",
        "pdf": "pdf",
        "jpg": "image",
        "jpeg": "image",
        "png": "image",
        "webp": "image",
        "gif": "image",
        "bmp": "image",
        "md": "markdown",
        "markdown": "markdown",
        "txt": "plain_text",
        "json": "json",
        "csv": "csv",
        "html": "html",
        "htm": "html",
        "xml": "xml",
        "yaml": "yaml",
        "yml": "yaml",
        "log": "log",
    }
    return mapping.get(ext, "plain_text")


def _count_words(text: str) -> int:
    """统计字数（中文字符 + 英文单词）。"""
    import re
    chinese = len(re.findall(r"[一-鿿]", text))
    english = len(re.findall(r"[a-zA-Z]+", text))
    return chinese + english


# ----------------------------------------------------------------
# 端点
# ----------------------------------------------------------------

@router.post("/upload", response_model=UploadResponse)
async def upload_file(file: UploadFile = File(...)):
    """
    上传文档/图片文件，自动解析为纯文本。

    支持格式：
    - 文本类：.md, .txt, .json, .csv, .html, .xml, .yaml, .log 等（本地读取）
    - Word 文档：.docx（后端 python-docx 解析）
    - PDF：.pdf（后端 pdfplumber 解析）
    - 图片：.jpg, .jpeg, .png, .webp, .gif, .bmp（Qwen VL 识别）
    """
    return await _parse_single_file(file)


# ---- 批量上传端点 ----

class BatchUploadItem(BaseModel):
    filename: str
    content: str
    format: str
    word_count: int
    error: str | None = None
    ok: bool = True


class BatchUploadResponse(BaseModel):
    files: list[BatchUploadItem]
    total_word_count: int
    merged_content: str


@router.post("/upload/batch", response_model=BatchUploadResponse)
async def upload_files_batch(files: list[UploadFile] = File(...)):
    """
    批量上传文件，自动解析每个文件并返回合并结果。

    - 每个文件独立解析，单个失败不影响其他文件
    - merged_content 是所有成功文件内容的拼接（用分隔线隔开）
    """
    if not files:
        raise HTTPException(status_code=400, detail="未提供任何文件")

    # 总大小检查
    total_bytes = 0
    max_total = int(os.getenv("UPLOAD_MAX_TOTAL_MB", "50")) * 1024 * 1024
    for f in files:
        # 读取文件以获取大小（FastAPI UploadFile 已缓存）
        pass

    results: list[BatchUploadItem] = []

    for file in files:
        total_bytes += file.size or 0

    if total_bytes > max_total:
        raise HTTPException(
            status_code=413,
            detail=f"文件总大小 ({total_bytes / 1024 / 1024:.1f}MB) 超过限制 ({max_total / 1024 / 1024:.0f}MB)",
        )

    for file in files:
        try:
            parsed = await _parse_single_file(file)
            results.append(BatchUploadItem(
                filename=parsed.filename,
                content=parsed.content,
                format=parsed.format,
                word_count=parsed.word_count,
                ok=True,
            ))
        except HTTPException as e:
            results.append(BatchUploadItem(
                filename=file.filename or "unknown",
                content="",
                format="",
                word_count=0,
                ok=False,
                error=e.detail if isinstance(e.detail, str) else str(e.detail),
            ))
        except Exception as e:
            results.append(BatchUploadItem(
                filename=file.filename or "unknown",
                content="",
                format="",
                word_count=0,
                ok=False,
                error=str(e),
            ))

    # 合并成功文件的内容
    ok_files = [r for r in results if r.ok and r.content.strip()]
    merged = "\n\n---\n\n".join(r.content for r in ok_files)
    total_words = sum(r.word_count for r in ok_files)

    return BatchUploadResponse(
        files=results,
        total_word_count=total_words,
        merged_content=merged,
    )


# ----------------------------------------------------------------
# 内部辅助
# ----------------------------------------------------------------

async def _parse_single_file(file: UploadFile) -> UploadResponse:
    """解析单个文件（供单文件和批量端点复用）。"""
    if not file.filename:
        raise HTTPException(status_code=400, detail="未提供文件名")

    ext = "." + file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALL_SUPPORTED:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式 '{ext}'，支持: {', '.join(sorted(ALL_SUPPORTED))}",
        )

    is_image = ext in IMAGE_EXTENSIONS

    file_bytes = await file.read()

    size_limit = MAX_IMAGE_SIZE if is_image else MAX_FILE_SIZE
    if len(file_bytes) > size_limit:
        raise HTTPException(
            status_code=413,
            detail=f"文件过大 ({len(file_bytes) / 1024 / 1024:.1f}MB)，上限为 {size_limit / 1024 / 1024:.0f}MB",
        )

    if len(file_bytes) == 0:
        raise HTTPException(status_code=400, detail="文件为空")

    try:
        if ext in IMAGE_EXTENSIONS:
            content = await _parse_image(file_bytes, file.filename or "image")
        elif ext == ".docx":
            content = _parse_docx(file_bytes)
        elif ext == ".pdf":
            content = _parse_pdf(file_bytes)
        else:
            content = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        raise HTTPException(
            status_code=400,
            detail="无法解码文件内容，请确保文件为 UTF-8 编码",
        )
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"文档解析失败: {str(e)}",
        )

    if not content.strip():
        raise HTTPException(status_code=400, detail="文档内容为空，无法提取文本")

    return UploadResponse(
        filename=file.filename,
        content=content,
        format=_detect_format(file.filename),
        word_count=_count_words(content),
    )
