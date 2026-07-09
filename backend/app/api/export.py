"""
多格式导出端点 —— 高质量 DOCX 和 PDF 生成。
DOCX: python-docx (with skills-guided formatting)
PDF: reportlab (professional document layout)
"""

import re
from io import BytesIO
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

# DOCX
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.units import cm, mm
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak,
    Table, TableStyle, ListFlowable, ListItem, KeepTogether,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus.flowables import HRFlowable

router = APIRouter(prefix="/api/export", tags=["export"])

EXPORT_FORMATS = ["docx", "pdf"]

# ============================================================
# Shared
# ============================================================
class ExportRequest(BaseModel):
    content: str = Field(..., min_length=1)
    format: str = Field(..., description="docx | pdf")


def _parse_md_blocks(markdown: str) -> list[dict]:
    """Parse Markdown into structured blocks for rendering."""
    blocks = []
    lines = markdown.strip().split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Code block
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1
            blocks.append({'type': 'code', 'text': '\n'.join(code_lines)})
            continue

        # Heading
        if line.startswith('# ') and not line.startswith('## '):
            blocks.append({'type': 'h1', 'text': line[2:].strip()})
        elif line.startswith('## '):
            blocks.append({'type': 'h2', 'text': line[3:].strip()})
        elif line.startswith('### '):
            blocks.append({'type': 'h3', 'text': line[4:].strip()})
        # Divider
        elif line.strip() == '---':
            blocks.append({'type': 'hr'})
        # Table
        elif line.strip().startswith('|') and line.strip().endswith('|'):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')])
                i += 1
            i -= 1
            # Filter separator row
            data = [r for r in table_rows if not all(re.match(r'^:?-{2,}:?$', c) for c in r)]
            if data:
                blocks.append({'type': 'table', 'rows': data})
        # Blockquote
        elif line.startswith('> '):
            quote_lines = []
            while i < len(lines) and lines[i].startswith('> '):
                quote_lines.append(lines[i][2:])
                i += 1
            i -= 1
            blocks.append({'type': 'quote', 'text': ' '.join(quote_lines)})
        # Unordered list item
        elif re.match(r'^\s*[-*]\s', line):
            items = []
            while i < len(lines) and re.match(r'^\s*[-*]\s', lines[i]):
                items.append(re.sub(r'^\s*[-*]\s', '', lines[i]))
                i += 1
            i -= 1
            blocks.append({'type': 'ul', 'items': items})
        # Ordered list item
        elif re.match(r'^\s*\d+\.\s', line):
            items = []
            while i < len(lines) and re.match(r'^\s*\d+\.\s', lines[i]):
                items.append(re.sub(r'^\s*\d+\.\s', '', lines[i]))
                i += 1
            i -= 1
            blocks.append({'type': 'ol', 'items': items})
        # Paragraph
        else:
            blocks.append({'type': 'p', 'text': line.strip()})
        i += 1
    return blocks


# ============================================================
# DOCX Export
# ============================================================
def _build_docx(blocks: list[dict]) -> BytesIO:
    doc = Document()

    # -- Page setup (A4) --
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # -- Style overrides --
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    # Set east-asian fallback font
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Microsoft YaHei"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Heading styles
    for i, (h_id, size, color) in enumerate([
        ('Heading 1', 22, '1E1B18'),
        ('Heading 2', 16, '3C4E5E'),
        ('Heading 3', 13, '4A6074'),
    ]):
        if h_id in [s.name for s in doc.styles]:
            hs = doc.styles[h_id]
        else:
            hs = doc.styles.add_style(h_id, WD_STYLE_TYPE.PARAGRAPH)
        hs.font.name = 'Arial'
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor.from_string(color)
        hs.paragraph_format.space_before = Pt(16 if i == 0 else 12)
        hs.paragraph_format.space_after = Pt(8)
        hs_rPr = hs.element.get_or_add_rPr()
        hs_rFonts = hs_rPr.find(qn('w:rFonts'))
        if hs_rFonts is None:
            hs_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Microsoft YaHei"/>')
            hs_rPr.append(hs_rFonts)
        else:
            hs_rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # -- Render blocks --
    for block in blocks:
        t = block['type']

        if t == 'h1':
            doc.add_heading(block['text'], level=1)
        elif t == 'h2':
            doc.add_heading(block['text'], level=2)
        elif t == 'h3':
            doc.add_heading(block['text'], level=3)
        elif t == 'hr':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
        elif t == 'quote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            run = p.add_run(block['text'])
            run.italic = True
            run.font.color.rgb = RGBColor(0x7D, 0x61, 0x32)
        elif t == 'ul':
            for item in block['items']:
                p = doc.add_paragraph(item, style='List Bullet')
        elif t == 'ol':
            for item in block['items']:
                p = doc.add_paragraph(item, style='List Number')
        elif t == 'table':
            rows = block['rows']
            if rows:
                ncols = max(len(r) for r in rows)
                # Pad rows to equal column count
                for r in rows:
                    while len(r) < ncols:
                        r.append('')
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = 'Light Grid Accent 1'
                for ri, row_data in enumerate(rows):
                    for ci, cell_text in enumerate(row_data):
                        cell = table.cell(ri, ci)
                        cell.text = cell_text
                        if ri == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
        elif t == 'code':
            for code_line in block['text'].split('\n'):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(code_line)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x3C, 0x4E, 0x5E)
        elif t == 'p':
            p = doc.add_paragraph()
            # Parse inline bold
            parts = re.split(r'(\*\*.*?\*\*)', block['text'])
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# PDF Export (reportlab)
# ============================================================
def _build_pdf(blocks: list[dict]) -> BytesIO:
    buffer = BytesIO()

    W, H = A4  # 595.27 x 841.89 points
    MARGIN = 2.54 * cm  # 1 inch

    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=MARGIN, rightMargin=MARGIN,
        topMargin=MARGIN, bottomMargin=MARGIN,
        title='Study Notes',
        author='Study Agent',
    )

    # Color palette
    INK = HexColor('#1E1B18')
    SOFT = HexColor('#5C5853')
    PRIMARY = HexColor('#3C4E5E')
    GOLD = HexColor('#B8944F')
    ACCENT_BG = HexColor('#FCF8F2')

    # Styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        'CNBody', parent=styles['Normal'],
        fontName='Helvetica', fontSize=10.5, leading=18,
        textColor=SOFT, spaceAfter=10,
    ))
    styles.add(ParagraphStyle(
        'CNH1', parent=styles['Heading1'],
        fontName='Helvetica-Bold', fontSize=20, leading=26,
        textColor=INK, spaceBefore=28, spaceAfter=12,
    ))
    styles.add(ParagraphStyle(
        'CNH2', parent=styles['Heading2'],
        fontName='Helvetica-Bold', fontSize=14, leading=20,
        textColor=PRIMARY, spaceBefore=20, spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        'CNH3', parent=styles['Heading3'],
        fontName='Helvetica-Bold', fontSize=11.5, leading=17,
        textColor=PRIMARY, spaceBefore=14, spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        'CNQuote', parent=styles['Normal'],
        fontName='Helvetica-Oblique', fontSize=10, leading=16,
        textColor=HexColor('#7D6132'), leftIndent=20,
        borderPadding=8, backColor=ACCENT_BG,
        spaceBefore=8, spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        'CNCode', parent=styles['Normal'],
        fontName='Courier', fontSize=8.5, leading=13,
        textColor=PRIMARY, leftIndent=12,
        spaceBefore=4, spaceAfter=4,
    ))

    story = []

    for block in blocks:
        t = block['type']

        if t == 'h1':
            story.append(Paragraph(block['text'], styles['CNH1']))
        elif t == 'h2':
            story.append(Paragraph(block['text'], styles['CNH2']))
        elif t == 'h3':
            story.append(Paragraph(block['text'], styles['CNH3']))
        elif t == 'hr':
            story.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#E8E4DE'), spaceBefore=10, spaceAfter=10))
        elif t == 'quote':
            story.append(Paragraph(block['text'], styles['CNQuote']))
        elif t == 'ul':
            items = [ListItem(Paragraph(it, styles['CNBody'])) for it in block['items']]
            story.append(ListFlowable(items, bulletType='bullet', bulletColor=INK, start='•', leftIndent=20, bulletFontSize=8))
        elif t == 'ol':
            items = [ListItem(Paragraph(it, styles['CNBody'])) for it in block['items']]
            story.append(ListFlowable(items, bulletType='1', bulletColor=INK, start='1', leftIndent=20))
        elif t == 'table':
            rows = block['rows']
            # Build table data
            table_data = []
            for ri, row in enumerate(rows):
                table_data.append([Paragraph(c, styles['CNBody']) for c in row])
            if table_data:
                col_w = (W - 2 * MARGIN) / max(len(r) for r in rows)
                t = Table(table_data, colWidths=[col_w] * max(len(r) for r in rows))
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), INK),
                    ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#FFFFFF')),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                    ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#E8E4DE')),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [HexColor('#FFFFFF'), HexColor('#F7F6F3')]),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('LEFTPADDING', (0, 0), (-1, -1), 10),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                ]))
                story.append(Spacer(1, 6))
                story.append(t)
                story.append(Spacer(1, 6))
        elif t == 'code':
            for code_line in block['text'].split('\n'):
                story.append(Paragraph(code_line.replace(' ', ' '), styles['CNCode']))
        elif t == 'p':
            # Parse inline bold
            parts = re.split(r'(\*\*.*?\*\*)', block['text'])
            para_parts = []
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    para_parts.append(f'<b>{part[2:-2]}</b>')
                else:
                    para_parts.append(part)
            story.append(Paragraph(''.join(para_parts), styles['CNBody']))

    doc.build(story)
    buffer.seek(0)
    return buffer


# ============================================================
# API Endpoint
# ============================================================
@router.post("")
async def export_document(request: ExportRequest):
    """
    将 Markdown 内容转换为指定格式并返回文件流。
    - format: `docx` (Word) 或 `pdf` (PDF)
    """
    if request.format not in EXPORT_FORMATS:
        raise HTTPException(400, detail=f"不支持的格式 '{request.format}'，可选: {EXPORT_FORMATS}")

    try:
        blocks = _parse_md_blocks(request.content)

        if request.format == 'docx':
            buffer = _build_docx(blocks)
            return StreamingResponse(
                buffer,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment; filename=notes.docx"},
            )
        else:  # pdf
            buffer = _build_pdf(blocks)
            return StreamingResponse(
                buffer,
                media_type="application/pdf",
                headers={"Content-Disposition": "attachment; filename=notes.pdf"},
            )
    except Exception as e:
        raise HTTPException(500, detail=f"导出失败: {str(e)}")
